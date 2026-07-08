import os
import tkinter as tk
from tkinter import filedialog, messagebox

import customtkinter as ctk
from PIL import Image, ImageTk
from docx import Document as DocxDocument
from ebooklib import epub

from pdfmagus.operations import format_converter
from pdfmagus.theme import COLORS, FONTS
from pdfmagus.widgets.buttons import primary_button, run_button


class FormatConversionPanel:
    def __init__(self, preview_canvas, log_event):
        self.preview_canvas = preview_canvas
        self.log_event = log_event

        self.convert_file_path = None
        self.convert_file_type = None

    def build(self, parent):
        ctk.CTkLabel(
            parent, text="Convert files to PDF", text_color=COLORS["primary"], font=("Arial", 12, "bold")
        ).pack(pady=(5, 15))

        file_frame = tk.Frame(parent, bg=COLORS["white"])
        file_frame.pack(fill="x", pady=10)

        primary_button(file_frame, "Select File", self.select_file_to_convert, width=150).pack(pady=5)

        self.lbl_file = ctk.CTkLabel(
            parent, text="No file selected", font=FONTS["label"], text_color=COLORS["black"]
        )
        self.lbl_file.pack(pady=5)

        tk.Frame(parent, bg="#eeeeee", height=2).pack(fill="x", pady=20)

        self.btn_convert = run_button(parent, "Convert to PDF", self.convert_file_to_pdf)
        self.btn_convert.pack(pady=10)
        self.btn_convert.configure(state="disabled")

        self.lbl_info = ctk.CTkLabel(
            parent,
            text="Supports: Word (.docx), Ebook (.epub), Images (.jpg, .png, .bmp)",
            font=FONTS["small"],
            text_color=COLORS["text_muted"],
        )
        self.lbl_info.pack(pady=10)

        if self.convert_file_path:
            file_name = os.path.basename(self.convert_file_path)
            self.lbl_file.configure(text=file_name)
            self.btn_convert.configure(state="normal")
            self.show_convert_preview()

    def select_file_to_convert(self):
        try:
            file_path = filedialog.askopenfilename(
                title="Select file to convert",
                filetypes=[
                    ("All supported", "*.docx *.epub *.jpg *.jpeg *.png *.bmp *.gif"),
                    ("Word documents", "*.docx"),
                    ("Ebooks", "*.epub"),
                    ("Images", "*.jpg *.jpeg *.png *.bmp *.gif"),
                    ("All files", "*.*"),
                ],
            )

            if not file_path or not os.path.exists(file_path):
                return

            self.convert_file_path = file_path
            file_ext = os.path.splitext(file_path)[1].lower()
            file_name = os.path.basename(file_path)

            if file_ext == ".docx":
                self.convert_file_type = "docx"
            elif file_ext == ".epub":
                self.convert_file_type = "epub"
            elif file_ext in (".jpg", ".jpeg", ".png", ".bmp", ".gif"):
                self.convert_file_type = "image"
            else:
                messagebox.showwarning("Unsupported type", "This file type is not supported.")
                return

            self.lbl_file.configure(text=file_name)
            self.btn_convert.configure(state="normal")
            self.log_event(f"File loaded for conversion: {file_name}")

            self.show_convert_preview()
        except Exception as e:
            self.log_event(f"Error selecting file: {e}")
            messagebox.showerror("Error", f"Error selecting file:\n{e}")

    def show_convert_preview(self):
        try:
            self.preview_canvas.delete("all")
            self.preview_canvas.image_list = []

            if not self.convert_file_path:
                return

            if self.convert_file_type == "docx":
                self._show_docx_preview()
            elif self.convert_file_type == "epub":
                self._show_epub_preview()
            elif self.convert_file_type == "image":
                self._show_image_preview()
        except Exception as e:
            self.preview_canvas.create_text(
                self.preview_canvas.winfo_width() // 2,
                self.preview_canvas.winfo_height() // 2,
                text=f"Error showing preview:\n{e}",
                fill="red",
            )

    def _show_docx_preview(self):
        try:
            doc = DocxDocument(self.convert_file_path)
            info_text = "WORD DOCUMENT\n"
            info_text += f"Paragraphs: {len(doc.paragraphs)}\n"
            info_text += f"Tables: {len(doc.tables)}\n\n"
            info_text += "Content (first 500 characters):\n"
            info_text += "-" * 50 + "\n"

            full_text = "\n".join(p.text for p in doc.paragraphs[:10])
            preview_text = full_text[:500] + ("..." if len(full_text) > 500 else "")
            info_text += preview_text

            self.lbl_info.configure(text=f"DOCX: {len(doc.paragraphs)} paragraphs")

            self.preview_canvas.create_text(
                10, 10, text=info_text, anchor="nw", font=("Consolas", 9), fill="black", justify="left"
            )
            self.preview_canvas.configure(scrollregion=self.preview_canvas.bbox("all"))
        except Exception as e:
            raise Exception(f"Error reading DOCX: {e}") from e

    def _show_epub_preview(self):
        try:
            book = epub.read_epub(self.convert_file_path)

            title = "No title"
            author = "No author"

            try:
                metadata = book.metadata
                if metadata:
                    titles = metadata.get("http://purl.org/dc/elements/1.1/", {}).get("title", [])
                    if titles:
                        title = list(titles)[0] if hasattr(titles, "__iter__") else titles

                    authors = metadata.get("http://purl.org/dc/elements/1.1/", {}).get("creator", [])
                    if authors:
                        author = list(authors)[0] if hasattr(authors, "__iter__") else authors
            except Exception:
                pass

            chapters = [item for item in book.get_items() if item.get_type() == 9]
            all_items = list(book.get_items())

            total_content = 0
            for item in all_items:
                try:
                    total_content += len(item.get_content())
                except Exception:
                    pass

            info_text = "EBOOK (EPUB)\n"
            info_text += f"Title: {title}\n"
            info_text += f"Author: {author}\n"
            info_text += f"Chapters: {len(chapters)}\n"
            info_text += f"Total items: {len(all_items)}\n"
            info_text += f"Content size: {total_content / 1024:.1f}KB"

            self.lbl_info.configure(text=f"EPUB: {len(chapters)} chapters")

            self.preview_canvas.create_text(
                10, 10, text=info_text, anchor="nw", font=("Consolas", 9), fill="black", justify="left"
            )
            self.preview_canvas.configure(scrollregion=self.preview_canvas.bbox("all"))
        except Exception as e:
            raise Exception(f"Error reading EPUB: {e}") from e

    def _show_image_preview(self):
        try:
            img = Image.open(self.convert_file_path)
            width, height = img.size

            preview_width = min(400, width)
            preview_height = min(300, height)
            img.thumbnail((preview_width, preview_height), Image.Resampling.LANCZOS)

            self.preview_photo = ImageTk.PhotoImage(img)
            self.preview_canvas.create_image(10, 10, anchor="nw", image=self.preview_photo)

            info_text = f"IMAGE\nDimensions: {width}x{height} pixels\nFormat: {img.format}"
            self.lbl_info.configure(text=info_text)
            self.preview_canvas.configure(scrollregion=self.preview_canvas.bbox("all"))
        except Exception as e:
            raise Exception(f"Error reading image: {e}") from e

    def convert_file_to_pdf(self):
        if not self.convert_file_path:
            messagebox.showwarning("Warning", "Please select a file first.")
            return

        try:
            output_path = filedialog.asksaveasfilename(
                title="Save converted PDF",
                defaultextension=".pdf",
                initialfile=os.path.splitext(os.path.basename(self.convert_file_path))[0] + ".pdf",
                filetypes=[("PDF files", "*.pdf")],
            )

            if not output_path:
                return

            self.btn_convert.configure(state="disabled")
            self.preview_canvas.create_text(
                self.preview_canvas.winfo_width() // 2,
                self.preview_canvas.winfo_height() // 2,
                text="Converting...",
                font=("Arial", 14),
                fill=COLORS["primary"],
            )
            self.preview_canvas.update()

            if self.convert_file_type == "docx":
                format_converter.docx_to_pdf(self.convert_file_path, output_path)
            elif self.convert_file_type == "epub":
                format_converter.epub_to_pdf(self.convert_file_path, output_path)
            elif self.convert_file_type == "image":
                format_converter.image_to_pdf(self.convert_file_path, output_path)

            self.log_event(f"File converted to PDF: {os.path.basename(output_path)}")
            messagebox.showinfo("Success", f"PDF converted successfully:\n{output_path}")
            self.show_convert_preview()
            self.btn_convert.configure(state="normal")
        except Exception as e:
            self.log_event(f"Error converting file: {e}")
            messagebox.showerror("Error", f"Error converting:\n{e}")
            self.show_convert_preview()
            self.btn_convert.configure(state="normal")
