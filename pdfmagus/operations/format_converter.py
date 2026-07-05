import os
import re

from PIL import Image
from docx import Document as DocxDocument
from ebooklib import epub
from reportlab.lib.pagesizes import letter
from reportlab.lib.utils import ImageReader
from reportlab.pdfgen import canvas


def wrap_text(text, canvas_obj, font_name, font_size, max_width):
    words = text.split()
    lines = []
    current_line = ""

    for word in words:
        test_line = (current_line + " " + word).strip()
        if canvas_obj.stringWidth(test_line, font_name, font_size) <= max_width:
            current_line = test_line
        else:
            if current_line:
                lines.append(current_line)
            current_line = word

    if current_line:
        lines.append(current_line)

    return lines


def docx_to_pdf(input_path, output_path):
    doc = DocxDocument(input_path)

    c = canvas.Canvas(output_path, pagesize=letter)
    width, height = letter
    margin = 50
    y_position = height - margin

    for paragraph in doc.paragraphs:
        if not paragraph.text.strip():
            continue

        font_size = 12
        font_name = "Helvetica"
        c.setFont(font_name, font_size)

        text = paragraph.text
        while text:
            if y_position < margin:
                c.showPage()
                y_position = height - margin

            words = text.split()
            line = ""
            remaining = ""

            for i, word in enumerate(words):
                test_line = (line + " " + word).strip()
                if c.stringWidth(test_line, font_name, font_size) < width - 2 * margin:
                    line = test_line
                else:
                    remaining = " ".join(words[i:])
                    break
            else:
                remaining = ""

            if line or not words:
                line = line or text
                c.drawString(margin, y_position, line)
                y_position -= font_size + 5
                text = remaining
            else:
                break

    for _ in doc.tables:
        if y_position < margin + 100:
            c.showPage()
            y_position = height - margin

        c.setFont("Helvetica", 10)
        c.drawString(margin, y_position, "[Table detected]")
        y_position -= 20

    c.save()


def epub_to_pdf(input_path, output_path):
    book = epub.read_epub(input_path)

    c = canvas.Canvas(output_path, pagesize=letter)
    width, height = letter
    margin = 50
    y_position = height - margin

    all_text = ""
    for item in book.get_items():
        if item.get_type() != 9:
            continue
        try:
            content = item.get_content().decode("utf-8", errors="ignore")
            content = re.sub(r"<script[^>]*>.*?</script>", "", content, flags=re.DOTALL)
            content = re.sub(r"<style[^>]*>.*?</style>", "", content, flags=re.DOTALL)
            content = re.sub(r"<[^>]+>", "\n", content)
            content = re.sub(r"\s+", " ", content)
            all_text += content + "\n\n"
        except Exception:
            pass

    if not all_text.strip():
        metadata = book.metadata
        if metadata:
            try:
                titles = metadata.get("http://purl.org/dc/elements/1.1/", {}).get("title", ["EPUB"])
                all_text = f"Title: {list(titles)[0]}\n\n"
            except Exception:
                all_text = "Converted EPUB\n\n"

    c.setFont("Helvetica", 11)

    for paragraph in all_text.split("\n"):
        paragraph = paragraph.strip()
        if not paragraph:
            y_position -= 10
            continue

        if y_position < margin + 20:
            c.showPage()
            y_position = height - margin

        wrapped_lines = wrap_text(paragraph, c, "Helvetica", 11, width - 2 * margin)
        for wrapped_line in wrapped_lines:
            if y_position < margin:
                c.showPage()
                y_position = height - margin

            c.drawString(margin, y_position, wrapped_line)
            y_position -= 14

    c.save()


def image_to_pdf(input_path, output_path):
    img = Image.open(input_path)

    if img.mode in ("RGBA", "LA", "P"):
        background = Image.new("RGB", img.size, (255, 255, 255))
        mask = img.split()[-1] if img.mode in ("RGBA", "LA") else None
        background.paste(img, mask=mask)
        img = background
    elif img.mode != "RGB":
        img = img.convert("RGB")

    page_width, page_height = letter
    img_width, img_height = img.size
    aspect_ratio = img_width / img_height

    margin = 40
    max_width = page_width - 2 * margin
    max_height = page_height - 2 * margin

    if aspect_ratio > max_width / max_height:
        new_width = int(max_width)
        new_height = int(new_width / aspect_ratio)
    else:
        new_height = int(max_height)
        new_width = int(new_height * aspect_ratio)

    img = img.resize((new_width, new_height), Image.Resampling.LANCZOS)

    temp_img_path = output_path.replace(".pdf", "_temp.png")
    img.save(temp_img_path)

    c = canvas.Canvas(output_path, pagesize=letter)
    img_reader = ImageReader(temp_img_path)

    x = (page_width - new_width) / 2
    y = (page_height - new_height) / 2
    c.drawImage(img_reader, x, y, width=new_width, height=new_height)
    c.save()

    if os.path.exists(temp_img_path):
        os.remove(temp_img_path)
